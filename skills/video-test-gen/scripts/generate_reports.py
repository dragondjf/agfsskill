#!/usr/bin/env python3
"""
生成测试用例 Excel 和 Word 报告

支持两种输入格式:
  1. analyze_frames.py 产出的 JSON（包含 test_cases 字段）
  2. 纯测试用例 JSON 数组

Excel 格式（与 excel_service 对齐）:
  ID | Title | Description | Preconditions | Priority | Step Number | Step Description | Expected Result
  同一用例的多步骤合并为同一组（首个步骤显示完整信息，后续步骤留空公共字段）

依赖: pip install pandas xlsxwriter python-docx
"""

import json
import os
import sys
import argparse


def _build_rows_from_test_cases(test_cases):
    """将测试用例列表扁平化为 Excel 行（与 excel_service.generate_excel 逻辑一致）"""
    rows = []
    for tc in test_cases:
        if not isinstance(tc, dict):
            continue

        tc_info = {
            "ID": tc.get("id", ""),
            "Title": tc.get("title", ""),
            "Description": tc.get("description", ""),
            "Preconditions": tc.get("preconditions", "") or "",
            "Priority": tc.get("priority", "") or "中",
        }

        steps = tc.get("steps", [])
        if not steps:
            # 没有步骤时仍占一行
            rows.append({
                **tc_info,
                "Step Number": "",
                "Step Description": "",
                "Expected Result": "",
            })
            continue

        for i, step in enumerate(steps):
            if i == 0:
                row = {
                    **tc_info,
                    "Step Number": step.get("step_number", i + 1),
                    "Step Description": step.get("description", ""),
                    "Expected Result": step.get("expected_result", ""),
                }
            else:
                row = {
                    "ID": "",
                    "Title": "",
                    "Description": "",
                    "Preconditions": "",
                    "Priority": "",
                    "Step Number": step.get("step_number", i + 1),
                    "Step Description": step.get("description", ""),
                    "Expected Result": step.get("expected_result", ""),
                }
            rows.append(row)

    return rows


def _get_priority_color(priority):
    """根据优先级返回背景色"""
    p = str(priority).strip().lower()
    if p in ("高", "high", "p0", "p1"):
        return "#FFC7CE"  # 浅红
    elif p in ("中", "medium", "p2"):
        return "#FFEB9C"  # 浅黄
    elif p in ("低", "low", "p3"):
        return "#C6EFCE"  # 浅绿
    return None


def save_to_excel(test_cases, excel_path):
    """
    将测试用例生成格式化的 Excel 文件（与 excel_service 对齐）

    格式特性:
    - 绿色背景加粗表头
    - 自动换行 + 顶部对齐
    - 优先级列根据级别着色
    - 多步骤用例的公共字段留空（视觉合并效果）
    - 合理列宽
    """
    import pandas as pd

    rows = _build_rows_from_test_cases(test_cases)

    if not rows:
        print("[警告] 无测试用例数据，生成空 Excel")
        rows = [{"ID": "", "Title": "", "Description": "", "Preconditions": "",
                 "Priority": "", "Step Number": "", "Step Description": "",
                 "Expected Result": ""}]

    df = pd.DataFrame(rows)
    columns = ["ID", "Title", "Description", "Preconditions", "Priority",
               "Step Number", "Step Description", "Expected Result"]

    os.makedirs(os.path.dirname(os.path.abspath(excel_path)), exist_ok=True)

    writer = pd.ExcelWriter(excel_path, engine="xlsxwriter")
    df.to_excel(writer, sheet_name="Test Cases", index=False)

    workbook = writer.book
    worksheet = writer.sheets["Test Cases"]

    # ===== 表头格式 =====
    header_format = workbook.add_format({
        "bold": True,
        "text_wrap": True,
        "valign": "top",
        "fg_color": "#D7E4BC",   # 浅绿
        "border": 1,
        "font_size": 11,
    })
    for col_num, value in enumerate(columns):
        worksheet.write(0, col_num, value, header_format)

    # ===== 列宽 =====
    col_widths = {
        "A": 12,   # ID
        "B": 30,   # Title
        "C": 40,   # Description
        "D": 30,   # Preconditions
        "E": 10,   # Priority
        "F": 12,   # Step Number
        "G": 45,   # Step Description
        "H": 45,   # Expected Result
    }
    for col_letter, width in col_widths.items():
        worksheet.set_column(f"{col_letter}:{col_letter}", width)

    # ===== 正文格式 =====
    normal_format = workbook.add_format({
        "text_wrap": True,
        "valign": "top",
        "border": 1,
    })

    # 预格式化的优先级颜色
    priority_formats = {}
    for p_name, color in [("高", "#FFC7CE"), ("中", "#FFEB9C"), ("低", "#C6EFCE")]:
        priority_formats[p_name] = workbook.add_format({
            "text_wrap": True,
            "valign": "top",
            "border": 1,
            "fg_color": color,
        })

    # 逐行写入（对 Priority 列着色）
    priority_col_idx = columns.index("Priority")
    for row_idx in range(1, len(rows) + 1):
        row_data = rows[row_idx - 1]
        for col_idx in range(len(columns)):
            col_name = columns[col_idx]
            value = row_data.get(col_name, "")

            # 对优先级列应用颜色
            if col_idx == priority_col_idx and value:
                fmt = priority_formats.get(str(value).strip(), normal_format)
            else:
                fmt = normal_format

            worksheet.write(row_idx, col_idx, value, fmt)

    # ===== 冻结首行 =====
    worksheet.freeze_panes(1, 0)

    writer.close()
    file_size = os.path.getsize(excel_path)
    print(f"Excel 已保存: {excel_path} ({len(rows)} 行, {file_size / 1024:.1f} KB)")


def save_to_word(test_cases, word_path):
    """
    将测试用例生成为图文并茂的 Word 报告

    特性:
    - 封面标题 + 用例统计
    - 每个用例一个章节，含元信息、来源截图、步骤表格
    - 缺失截图时优雅降级
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("测试用例报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 统计信息
    total = len(test_cases)
    from collections import Counter
    priorities = Counter(tc.get("priority", "未知") for tc in test_cases)

    stats_para = doc.add_paragraph()
    stats_para.add_run(f"共 {total} 个测试用例").bold = True
    stats_para.add_run(f"\n优先级分布: " + ", ".join(
        f"{k}: {v}" for k, v in sorted(priorities.items())
    ))

    doc.add_paragraph()  # 空行

    for tc in test_cases:
        tc_id = tc.get("id", "TC-???")
        tc_title = tc.get("title", "未命名用例")

        doc.add_heading(f"{tc_id}: {tc_title}", level=1)

        # 元信息表格
        meta = doc.add_table(rows=3, cols=2, style="Table Grid")
        meta.rows[0].cells[0].text = "优先级"
        meta.rows[0].cells[1].text = str(tc.get("priority", "中"))
        meta.rows[1].cells[0].text = "描述"
        meta.rows[1].cells[1].text = str(tc.get("description", ""))
        meta.rows[2].cells[0].text = "前置条件"
        meta.rows[2].cells[1].text = str(tc.get("preconditions", ""))

        # 插入来源图片（如果有）
        img_path = tc.get("_source_image", "")
        if img_path and os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(5))
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败: {e}]")

        # 测试步骤表格
        steps = tc.get("steps", [])
        if steps:
            doc.add_heading("测试步骤", level=2)
            table = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "#"
            hdr_cells[1].text = "步骤描述"
            hdr_cells[2].text = "预期结果"

            for step in steps:
                row_cells = table.add_row().cells
                row_cells[0].text = str(step.get("step_number", ""))
                row_cells[1].text = str(step.get("description", ""))
                row_cells[2].text = str(step.get("expected_result", ""))

        doc.add_page_break()

    doc.save(word_path)
    file_size = os.path.getsize(word_path)
    print(f"Word 报告已保存: {word_path} ({file_size / 1024:.1f} KB)")


def main():
    parser = argparse.ArgumentParser(description="根据分析结果生成 Excel 和 Word 报告")
    parser.add_argument("analysis_json", help="分析结果 JSON 文件路径")
    parser.add_argument("--excel", default=None, help="Excel 输出路径")
    parser.add_argument("--word", default=None, help="Word 输出路径")
    args = parser.parse_args()

    if not args.excel and not args.word:
        print("[错误] 请至少指定 --excel 或 --word 其中一个输出")
        sys.exit(1)

    with open(args.analysis_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    # 兼容两种格式: {"test_cases": [...]} 或直接 [...]
    if isinstance(data, dict):
        test_cases = data.get("test_cases", [])
    elif isinstance(data, list):
        test_cases = data
    else:
        print("[错误] JSON 格式不支持，期望 dict 或 list")
        sys.exit(1)

    print(f"加载了 {len(test_cases)} 个测试用例")

    if args.excel:
        save_to_excel(test_cases, args.excel)
    if args.word:
        save_to_word(test_cases, args.word)
