#!/usr/bin/env python3
"""
测试记录生成器 - 基于 YFJZ-R805-05 模板格式生成软件测试记录

打开参考模板 docx，保留封面表和修改记录表，仅替换附录3内容。
每个测试用例生成 16x8 表格（纯 XML 构建，绕过 python-docx cell.merge 陷阱），
并嵌入对应关键帧截图。

文档结构（从模板继承）:
  封面表(7x2无边框) → 文档修改记录(5x5) → 附录3 测试记录
    → 按测试类型分组(功能测试/边界测试/...)
      → Heading2(TC编号) → 16x8表格 → 居中图片(5.51in) → 2空段落

参考: video-test-gen SKILL.md 步骤5

用法:
    python generate_test_records.py <analysis.json> -o <output.docx>
    python generate_test_records.py <analysis.json> -o <output.docx> --keyframes-dir ./keyframes_std

依赖: pip install python-docx Pillow
"""

import json
import os
import sys
import argparse
from collections import OrderedDict

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==================== XML 单元格构建 ====================

def mk_cell(text, grid_span=1, shading=None, bold=False, font="宋体", fsize="18",
            alignment="left"):
    """
    构建单个 w:tc XML 元素（纯 XML，不依赖 python-docx cell 缓存）。

    Args:
        text: 单元格文本，必须为 str（传入 int 会报类型错误）
        grid_span: 水平合并列数（默认 1，不合并）
        shading: 底纹颜色，如 "D9E2F3"（None 表示无底纹）
        bold: 是否加粗
        font: 字体名称（默认宋体）
        fsize: 字号半磅值，"18"=9pt, "24"=12pt
        alignment: 段落对齐（"left" / "center"）

    Returns:
        OxmlElement (w:tc)
    """
    text = str(text) if text is not None else ""

    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')

    # gridSpan
    if grid_span > 1:
        gs = OxmlElement('w:gridSpan')
        gs.set(qn('w:val'), str(grid_span))
        tcPr.append(gs)

    # 垂直居中
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)

    # 底纹
    if shading:
        sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear')
        sh.set(qn('w:color'), 'auto')
        sh.set(qn('w:fill'), shading)
        tcPr.append(sh)

    tc.append(tcPr)

    # 段落
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment)
    pPr.append(jc)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '0')
    sp.set(qn('w:line'), '280')
    sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)
    p.append(pPr)

    # Run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    for sz_tag in ('w:sz', 'w:szCs'):
        sz = OxmlElement(sz_tag)
        sz.set(qn('w:val'), fsize)
        rPr.append(sz)
    fonts = OxmlElement('w:rFonts')
    for fk in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
        fonts.set(qn(fk), font)
    rPr.append(fonts)
    r.append(rPr)

    t_elem = OxmlElement('w:t')
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_elem.text = text
    r.append(t_elem)
    p.append(r)
    tc.append(p)

    return tc


# ==================== 16x8 TC 表格构建 ====================

def build_tc_table(tc, frame_num, timestamp):
    """
    构建单个测试用例的 16 行 x 8 列表格（纯 XML）。

    表格结构:
      Row 0:  [4+4] 用例名称 | 用例标识
      Row 1:  [8]   用例描述
      Row 2:  [8]   用例输入（优先级+描述）
      Row 3:  [8]   测试类型
      Row 4:  [8]   前提和约束（帧信息）
      Row 5:  [8]   测试终止条件
      Row 6:  [8]   测试过程
      Row 7:  [1+2+2+2+1] 步骤表头
      Row 8-13: [1+2+2+2+1] 步骤数据行（最多6行）
      Row 14: [8]   测试结论
      Row 15: [4+4] 测试人员 | 测试日期

    Args:
        tc: 测试用例字典
        frame_num: 关键帧编号（从1开始）
        timestamp: 关键帧时间戳（秒）

    Returns:
        OxmlElement (w:tbl)
    """
    SHADING = "D9E2F3"

    tbl = OxmlElement('w:tbl')

    # tblPr
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    borders = OxmlElement('w:tblBorders')
    for bname in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        b.set(qn('w:space'), '0')
        borders.append(b)
    tblPr.append(borders)
    tbl.append(tblPr)

    # tblGrid
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(8):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), '1080')
        tblGrid.append(gc)
    tbl.append(tblGrid)

    # 用例数据
    tc_id = tc.get("id", "")
    title = tc.get("title", "")
    description = tc.get("description", "")
    priority = tc.get("priority", "中")
    test_type = tc.get("test_type", "功能测试")
    steps = tc.get("steps", [])

    def add_label_row(content, span=8):
        tr = OxmlElement('w:tr')
        tr.append(mk_cell(content, grid_span=span, shading=SHADING, bold=True))
        tbl.append(tr)

    # Row 0: 用例名称 + 用例标识
    tr = OxmlElement('w:tr')
    tr.append(mk_cell(f"测试用例名称：{title}", grid_span=4, shading=SHADING, bold=True))
    tr.append(mk_cell(f"测试用例标识：{tc_id}", grid_span=4, shading=SHADING, bold=True))
    tbl.append(tr)

    # Row 1-6: 标签行
    add_label_row(f"测试用例描述：{description}")
    add_label_row(f"测试用例输入：优先级：{priority}。{description}")
    add_label_row(f"测试类型：{test_type}")
    add_label_row(f"前提和约束：视频第 {frame_num} 帧 ({timestamp}s)")
    add_label_row(
        "测试终止条件：正常终止：该测试项的所有测试用例都正常终止。"
        "异常终止：测试过程中出现异常情况，需记录异常原因并终止测试。"
    )
    add_label_row("测试过程")

    # Row 7: 步骤表头
    tr = OxmlElement('w:tr')
    for label, span in [("序号", 1), ("输入及操作步骤", 2),
                         ("期望测试结果", 2), ("评估准则", 2),
                         ("实际测试结果", 1)]:
        tr.append(mk_cell(label, grid_span=span, shading=SHADING, bold=True))
    tbl.append(tr)

    # Row 8-13: 步骤数据行（最多6行）
    for i in range(6):
        tr = OxmlElement('w:tr')
        if i < len(steps):
            step = steps[i]
            sn = str(step.get("step_number", i + 1))
            sd = step.get("description", "")
            er = step.get("expected_result", "")
            tr.append(mk_cell(sn, grid_span=1))
            tr.append(mk_cell(sd, grid_span=2))
            tr.append(mk_cell(er, grid_span=2))
            tr.append(mk_cell("与预期结果一致", grid_span=2))
            tr.append(mk_cell("", grid_span=1))
        # 超出实际步骤数的行不添加任何单元格，保持空行
        tbl.append(tr)

    # Row 14: 测试结论
    add_label_row("测试结论")

    # Row 15: 测试人员 + 测试日期
    tr = OxmlElement('w:tr')
    tr.append(mk_cell("测试人员", grid_span=4, shading=SHADING, bold=True))
    tr.append(mk_cell("测试日期", grid_span=4, shading=SHADING, bold=True))
    tbl.append(tr)

    return tbl


# ==================== 文档构建 ====================

def _safe_tag(elem):
    """安全提取 XML 标签名（兼容多命名空间）。"""
    return elem.tag.rsplit('}', 1)[-1] if '}' in elem.tag else elem.tag


def _remove_appendix3(body):
    """删除 body 中 '附录3 测试记录' 段落及之后的所有元素。"""
    to_remove = []
    found = False
    for elem in body:
        tag = _safe_tag(elem)
        if tag == 'p':
            text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
            if '附录3 测试记录' in text:
                found = True
                to_remove.append(elem)
                continue
        if found:
            to_remove.append(elem)
    for e in to_remove:
        body.remove(e)
    return len(to_remove)


def _insert_after(body, ref_elem, new_elem):
    """将 new_elem 插入到 ref_elem 之后。"""
    ref_idx = list(body).index(ref_elem)
    body.insert(ref_idx + 1, new_elem)
    return ref_idx + 1


def _add_image_paragraph(doc, body, ref_elem, img_path, width_inches=5.51):
    """
    创建居中图片段落并插入到 ref_elem 之后。

    使用 add_paragraph 创建（添加到 body 末尾），然后 remove + insert 到目标位置。

    Args:
        doc: Document 对象
        body: doc.element.body
        ref_elem: 图片应紧跟在其后的元素
        img_path: 图片文件路径
        width_inches: 图片宽度（英寸）

    Returns:
        插入后的图片段落元素
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    p_elem = p._element
    body.remove(p_elem)
    idx = _insert_after(body, ref_elem, p_elem)
    return p_elem


def _add_empty_paragraphs(doc, body, ref_elem, count=2):
    """
    创建指定数量的空段落，依次插入到 ref_elem 之后。

    Returns:
        最后一个空段落元素
    """
    prev = ref_elem
    for _ in range(count):
        ep = doc.add_paragraph()._element
        body.remove(ep)
        prev_idx = list(body).index(prev)
        body.insert(prev_idx + 1, ep)
        prev = ep
    return prev


def generate_test_records(test_cases, output_path, template_path=None,
                          keyframes_dir=None, frame_info=None):
    """
    基于 YFJZ-R805-05 模板生成软件测试记录。

    打开参考模板 docx，保留封面表和修改记录表，仅替换附录3内容。
    每个用例按 Heading2 → 16x8 表格 → 居中图片 → 2 空段落的顺序排列。

    Args:
        test_cases: 测试用例列表（来自 analysis_merged.json）
        output_path: 输出 docx 路径
        template_path: 模板 docx 路径（默认使用 skill 内置模板）
        keyframes_dir: 关键帧标准 JPEG 目录（用于嵌入图片）
        frame_info: 帧信息列表 [[path, timestamp], ...]（默认从 keyframes_dir 扫描）
    """
    # ---- 默认模板路径 ----
    if template_path is None:
        skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(skill_dir, "references", "test_records_template.docx")

    if not os.path.exists(template_path):
        print(f"[错误] 模板文件不存在: {template_path}")
        sys.exit(1)

    # ---- 帧信息 ----
    num_frames = 0
    img_files = []

    if keyframes_dir and os.path.isdir(keyframes_dir):
        img_files = sorted([f for f in os.listdir(keyframes_dir)
                            if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
        num_frames = len(img_files)

    if frame_info is None:
        if keyframes_dir and img_files:
            # 从文件名推断时间戳（备用方案）
            frame_info = []
            for i, f in enumerate(img_files):
                # 尝试从文件名提取时间戳，失败则用索引*8
                ts = i * 8.0
                frame_info.append([os.path.join(keyframes_dir, f), ts])
        else:
            frame_info = []

    num_frames = max(num_frames, len(frame_info))

    print(f"模板: {template_path}")
    print(f"用例数: {len(test_cases)}")
    print(f"关键帧数: {num_frames}")

    # ---- 打开模板 ----
    doc = Document(template_path)
    body = doc.element.body

    # ---- 删除附录3及之后 ----
    removed = _remove_appendix3(body)
    print(f"删除旧内容: {removed} 个元素")

    # ---- 添加附录3标题 ----
    doc.add_heading('附录3 测试记录', level=1)

    # ---- 按测试类型分组 ----
    type_order = ["功能测试", "边界测试", "安全性测试", "性能测试"]
    groups = OrderedDict()
    for tc in test_cases:
        t = tc.get("test_type", "功能测试")
        groups.setdefault(t, []).append(tc)

    sorted_groups = OrderedDict()
    for t in type_order:
        if t in groups:
            sorted_groups[t] = groups[t]
    for t, cases in groups.items():
        if t not in sorted_groups:
            sorted_groups[t] = cases

    # ---- 遍历分组生成 TC ----
    tc_idx = 0
    for group_name, cases in sorted_groups.items():
        print(f"\n  分组: {group_name} ({len(cases)} 个用例)")
        doc.add_heading(group_name, level=1)

        for tc in cases:
            tc_id = tc.get("id", f"TC-{tc_idx + 1:03d}")

            # Heading2
            h2 = doc.add_heading(tc_id, level=2)
            h2_elem = h2._element

            # 帧映射
            frame_idx = tc_idx % max(num_frames, 1)
            frame_num = frame_idx + 1
            timestamp = frame_info[frame_idx][1] if frame_idx < len(frame_info) else 0

            # 16x8 表格
            tbl_elem = build_tc_table(tc, frame_num, timestamp)
            _insert_after(body, h2_elem, tbl_elem)

            # 图片
            if img_files:
                frame_path = os.path.join(keyframes_dir, img_files[frame_idx])
                if os.path.exists(frame_path):
                    p_img = _add_image_paragraph(doc, body, tbl_elem, frame_path)
                    # 2 个空段落
                    _add_empty_paragraphs(doc, body, p_img, count=2)
                else:
                    print(f"    [警告] 图片不存在: {frame_path}")
            else:
                # 无图片时插入 2 个空段落
                _add_empty_paragraphs(doc, body, tbl_elem, count=2)

            tc_idx += 1

    # ---- 保存（文件占用容错） ----
    try:
        doc.save(output_path)
        print(f"\n已保存: {output_path}")
        final_path = output_path
    except PermissionError:
        base, ext = os.path.splitext(output_path)
        alt_path = f"{base}_v2{ext}"
        doc.save(alt_path)
        print(f"\n文件被占用，已保存: {alt_path}")
        final_path = alt_path

    file_size = os.path.getsize(final_path)
    print(f"文件大小: {file_size / 1024:.1f} KB")


# ==================== CLI ====================

def main():
    parser = argparse.ArgumentParser(
        description="基于 YFJZ-R805-05 模板生成软件测试记录",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基本用法（使用内置模板）
  python generate_test_records.py analysis_merged.json -o test_records.docx

  # 指定关键帧目录
  python generate_test_records.py analysis_merged.json -o test_records.docx \\
      --keyframes-dir ./keyframes_std

  # 指定自定义模板
  python generate_test_records.py analysis_merged.json -o test_records.docx \\
      --template ./references/test_records_template.docx
        """
    )
    parser.add_argument("input_json", help="测试用例 JSON 文件路径（analysis_merged.json）")
    parser.add_argument("-o", "--output", required=True, help="输出 docx 文件路径")
    parser.add_argument("-t", "--template", default=None,
                        help="模板 docx 路径（默认使用 skill 内置模板）")
    parser.add_argument("-k", "--keyframes-dir", default=None,
                        help="关键帧标准 JPEG 目录（用于嵌入截图）")

    args = parser.parse_args()

    # 加载 JSON
    with open(args.input_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    if isinstance(data, dict):
        test_cases = data.get("test_cases", [])
    elif isinstance(data, list):
        test_cases = data
    else:
        print("[错误] JSON 格式不支持，需要 dict（含 test_cases）或 list")
        sys.exit(1)

    generate_test_records(
        test_cases=test_cases,
        output_path=args.output,
        template_path=args.template,
        keyframes_dir=args.keyframes_dir,
    )


if __name__ == "__main__":
    main()
